import streamlit as st
import pandas as pd
import statsmodels.api as sm
import requests
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from docx import Document
from textwrap3 import wrap
import tempfile

# サンプルデータの作成
data_mercedes = {
    'ブランドイメージ': [4.5, 4.6, 4.7, 4.8, 4.9, 5.0, 5.1, 5.2, 5.3, 5.4, 5.5, 5.6, 5.7, 5.8, 5.9, 6.0, 6.1, 6.2, 6.3, 6.4, 6.5, 6.6, 6.7, 6.8, 6.9, 4.0, 4.1, 4.2, 4.3, 4.4, 4.5, 4.6, 4.7, 4.8, 4.9, 5.0, 5.1, 5.2, 5.3, 5.4, 5.5, 5.6, 5.7, 5.8, 5.9, 6.0, 6.1, 6.2, 6.3, 6.4,
    6.5, 6.6, 6.7, 6.8, 6.9, 4.0, 4.1, 4.2, 4.3, 4.4,
    4.5, 4.6, 4.7, 4.8, 4.9, 5.0, 5.1, 5.2, 5.3, 5.4,
    5.5, 5.6, 5.7, 5.8, 5.9, 6.0, 6.1, 6.2, 6.3, 6.4,
    6.5, 6.6, 6.7, 6.8, 6.9, 4.0, 4.1, 4.2, 4.3, 4.4,
    4.5, 4.6, 4.7, 4.8, 4.9, 5.0, 5.1, 5.2, 5.3, 5.4],
    '顧客体験': [5, 4, 5, 4.5, 5, 4.1, 5.1, 4.2, 5.2, 4.3, 5.3, 4.4, 5.4, 4.5, 5.5, 4.6, 5.6, 4.7, 5.7, 4.8, 5.8, 4.9, 5.9, 4.0, 5.0, 4.1, 5.1, 4.2, 5.2, 4.3, 5.3, 4.4, 5.4, 4.5, 5.5, 4.6, 5.6, 4.7, 5.7, 4.8, 5.8, 4.9, 5.9, 4.0, 5.0, 4.1, 5.1, 4.2, 5.2, 4.3, 5.3, 4.4, 5.4, 4.5, 5.5, 4.6, 5.6, 4.7, 5.7, 4.8, 5.8, 4.9, 5.9, 4.0, 5.0, 4.1, 5.1, 4.2, 5.2, 4.3, 5.3, 4.4, 5.4, 4.5, 5.5, 4.6, 5.6, 4.7, 5.7, 4.8, 5.8, 4.9, 5.9, 4.0, 5.0, 4.1, 5.1, 4.2, 5.2, 4.3, 5.3, 4.4, 5.4, 4.5, 5.5, 4.6, 5.6, 4.7, 5.7, 4.8],
    'リレーションシップ営業': [4.2, 4.0, 4.8, 4.5, 4.6] * 20,
    'ローン・リースプラン': [3.5, 4.0, 3.8, 3.7, 4.1] * 20,
    '保証サービス': [4.5, 4.6, 4.4, 4.3, 4.7] * 20,
    '広告費': [60, 70, 50, 40, 80] * 20,
    '価格設定': [800, 850, 820, 830, 810] * 20,
    '売上': [500, 480, 510, 495, 530] * 20
}

data_bmw = {
    'ブランドイメージ': [4.5, 4.6, 4.7, 4.8, 4.9] * 20,
    '顧客体験': [4.8, 4.7, 4.9, 4.6, 5.0] * 20,
    'リレーションシップ営業': [4.1, 4.3, 4.2, 4.4, 4.5] * 20,
    'ローン・リースプラン': [3.8, 3.9, 4.0, 4.1, 4.2] * 20,
    '保証サービス': [4.3, 4.4, 4.5, 4.6, 4.7] * 20,
    '広告費': [55, 65, 75, 85, 95] * 20,
    '価格設定': [780, 790, 800, 810, 820] * 20,
    '売上': [490, 500, 510, 520, 530] * 20
}

data_nextage = {
    '価格': [200, 150, 300, 250, 100] * 20,
    '広告費': [50, 40, 60, 70, 30] * 20,
    '在庫数': [100, 120, 80, 90, 110] * 20,
    '車両品質': [4.5, 4.0, 4.8, 4.6, 4.2] * 20,
    '店舗立地': [3, 5, 4, 2, 4] * 20,
    '成約率': [0.5, 0.6, 0.55, 0.65, 0.6] * 20,
    '売上': [300, 250, 350, 330, 210] * 20
}

# フォントをWebからダウンロードして登録
font_url = 'https://www.ryhintl.com/font-nasu/Nasu-Regular.ttf'
font_name = 'MS-Gothic'
response = requests.get(font_url)
with tempfile.NamedTemporaryFile(delete=False, suffix=".ttf") as tf:
    tf.write(response.content)
pdfmetrics.registerFont(TTFont(font_name, tf.name))

# Word生成関数
def generate_word(content):
    doc = Document()
    doc.add_heading("Regression Result Analysis", level=1)
    for line in content.split("\n"):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# PDF生成関数
def generate_pdf(content):
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    p.setFont(font_name, 8)
    p.drawString(100, 800, "Regression Result Analysis")
    text = p.beginText(100, 780)
    text.setFont(font_name, 10)
    text.setLeading(14)
    for line in content.split("\n"):
        wrapped_lines = wrap(line, 50)
        for wrapped_line in wrapped_lines:
            text.textLine(wrapped_line)
    #for line in content.split("\n"):
        #text.textLine(line)
    p.drawText(text)
    p.showPage()
    p.save()
    buffer.seek(0)
    return buffer

# Streamlitアプリの設定
st.title('セールス分析')
st.write('セールスデータの回帰分析（因果探索）')

# データセットの選択
dataset = st.selectbox('データセット', ('Mercedes', 'BMW', 'NEXSTAGE'))

# データフレームの選択
if dataset in ['Mercedes', 'BMW']:
    dataframe = st.selectbox('データフレーム', ('ブランド', 'サービス'))
else:
    dataframe = st.selectbox('データフレーム', ('車両', '価格'))

# データフレームの作成
if dataset == 'Mercedes':
    df = pd.DataFrame(data_mercedes)
elif dataset == 'BMW':
    df = pd.DataFrame(data_bmw)
else:
    df = pd.DataFrame(data_nextage)

# 説明変数と目的変数に分ける
if dataframe == 'サービス':
    X = df[['ローン・リースプラン', '保証サービス', '広告費']]
    y = df['売上']
elif dataframe == 'ブランド':
    X = df[['ブランドイメージ', '顧客体験', 'リレーションシップ営業']]
    y = df['売上']
elif dataframe == '車両':
    X = df[['車両品質', '店舗立地', '成約率']]
    y = df['売上']
else:
    X = df[['価格', '広告費', '在庫数']]
    y = df['売上']

# 定数項を追加
X = sm.add_constant(X)

# 重回帰分析モデルの作成とフィッティング
model = sm.OLS(y, X).fit()
results_summary = model.summary()

# Groq API設定
API_URL = 'https://api.groq.com/openai/v1/'
MODEL = 'Llama-3.1-70b-Versatile'
API_KEY = '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'
maxTokens = 4096
headers = {
    'Content-Type': 'application/json',
    'Authorization': f'Bearer {API_KEY}'
}

# 分析ボタンの作成
if st.button('分析'):
    data = {
        'model': MODEL,
        'max_tokens': maxTokens,
        'messages': [
            {
                'role': 'system',
                'content': '貴方は専門家です。できるだけわかりやすく答えてください。必ず、日本語で答えてください。'
            },
            {
                'role': 'user',
                'content': '以下の文章を分析してください。' + str(results_summary)
            }
        ]
    }
    
    response = requests.post(f'{API_URL}chat/completions', headers=headers, json=data)
    groqResp = response.json()['choices'][0]['message']['content']

    # 結果をsession_stateに保存
    st.session_state['results_summary'] = results_summary
    st.session_state['groqResp'] = groqResp
    st.session_state['contents'] = str(results_summary) + "\n\n" + groqResp

# 分析結果の表示
if 'results_summary' in st.session_state and 'groqResp' in st.session_state:
    st.subheader('Regression 結果')
    st.text(st.session_state['results_summary'])
    st.subheader('結果分析')
    st.text_area('Result Analysis', st.session_state['groqResp'], height=300)

# ダウンロードボタン
if 'contents' in st.session_state:
    word_buffer = generate_word(st.session_state['contents'])
    st.download_button("Download Word", word_buffer, "分析結果.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    pdf_buffer = generate_pdf(st.session_state['contents'])
    st.download_button("Download PDF", pdf_buffer, "分析結果.pdf", "application/pdf")    
